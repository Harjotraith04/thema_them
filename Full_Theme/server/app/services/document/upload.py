"""
Document upload and file processing service
"""
from sqlalchemy.orm import Session
from typing import Optional, Dict, Any
import cloudinary
import cloudinary.uploader
import hashlib
import pypdf
import pandas as pd
from docx import Document as DocxDocument
import io
import pathlib
from app.schemas.document import DocumentUpload
# from app.schemas.document_segment import DocumentSegmentOut
from app.core.permissions import PermissionChecker
from app.core.config import settings
from app.models.document import Document, DocumentType
# from app.models.document_segment import DocumentSegment
from app.models.user import User

cloudinary.config(
    cloud_name=settings.CLOUDINARY_CLOUD_NAME,
    api_key=settings.CLOUDINARY_API_KEY,
    api_secret=settings.CLOUDINARY_API_SECRET
)


class DocumentUploadService:

    @staticmethod
    def create_document(
        db: Session,
        name: str,
        description: Optional[str],
        document_type: DocumentType,
        project_id: int,
        uploaded_by_id: int,
        file_content: bytes,
        filename: str
    ) -> DocumentUpload:
        """Create a new document with file processing"""

        # Get user object
        user = db.query(User).filter(User.id == uploaded_by_id).first()
        if not user:
            raise ValueError("User not found")

        # Check user access to project
        project = PermissionChecker.check_project_access(
            db, project_id, user, raise_exception=False
        )
        if not project:
            raise ValueError("Project not found or access denied")

        file_size = len(file_content)
        file_hash = hashlib.sha256(file_content).hexdigest()

        # Upload to Cloudinary
        upload_result = cloudinary.uploader.upload(
            file_content,
            resource_type="raw",
            public_id=f"documents/{project_id}/{file_hash}",
            use_filename=True,
            unique_filename=False
        )

        cloudinary_public_id = upload_result["public_id"]
        cloudinary_url = upload_result["secure_url"]
        # Extract segments and metadata based on file type and filename
        segments_data, file_metadata = DocumentUploadService._extract_segments(
            file_content, document_type, filename
        )

        extracted_content = DocumentUploadService._extract_full_content(
            file_content, document_type, filename
        )

        # Create the document record
        document = Document(
            name=name,
            description=description,
            document_type=document_type,
            content=extracted_content,
            file_size=file_size,
            file_hash=file_hash,
            cloudinary_public_id=cloudinary_public_id,
            cloudinary_url=cloudinary_url,
            file_metadata=file_metadata,
            project_id=project_id,
            uploaded_by_id=uploaded_by_id
        )
        db.add(document)
        db.commit()
        db.refresh(document)

        # Create DocumentSegment records from the extracted content
        # print(f"Document created: {document.name} (ID: {document.id})")
        # if segments_data and "segments" in segments_data:
        #     DocumentUploadService._create_document_segments(
        #         db, document, segments_data["segments"])
        #     print(f"Created {len(segments_data['segments'])} segments for document {document.id}")
        #     db.refresh(document)

        uploaded_doc = DocumentUpload(
            id=int(getattr(document, "id")),
            name=str(document.name),
            cloudinary_url=str(document.cloudinary_url),
            content=str(getattr(document, "content", "")),
            file_size=int(getattr(document, "file_size", 0)),
            upload_status="success"
        )
        return uploaded_doc

    @staticmethod
    def _extract_full_content(file_content: bytes, document_type: DocumentType, filename: str) -> str:
        """Extract full text content for the document content column"""
        try:
            if document_type == DocumentType.TEXT:
                clean_content = file_content.replace(b'\x00', b'')
                return clean_content.decode('utf-8', errors='replace')

            elif document_type == DocumentType.PDF:
                pdf_file = io.BytesIO(file_content)
                pdf_reader = pypdf.PdfReader(pdf_file)
                content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    content += f"--- Page {page_num + 1} ---\n"
                    content += page_text + "\n\n"
                return content

            elif document_type == DocumentType.DOCX:
                doc_file = io.BytesIO(file_content)
                doc = DocxDocument(doc_file)
                content = ""
                for paragraph in doc.paragraphs:
                    para_text = paragraph.text
                    if para_text or len(para_text.strip()) == 0:
                        content += para_text + "\n"
                    if not para_text.strip():
                        content += "\n"
                return content

            elif document_type == DocumentType.CSV:
                ext = pathlib.Path(filename or "").suffix.lower()
                if ext in (".xlsx", ".xls"):
                    try:
                        excel_buffer = io.BytesIO(file_content)
                        df = pd.read_excel(excel_buffer)
                        content = f"Excel file with {len(df)} rows and {len(df.columns)} columns\n"
                        content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
                        for row_index, (_, row) in enumerate(df.iterrows()):
                            row_values = [
                                f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])]
                            if row_values:
                                content += f"Row {row_index + 1}: " + \
                                    " | ".join(row_values) + "\n"
                        return content
                    except Exception:
                        return "[Error extracting Excel content]"
                else:
                    clean_content = file_content.replace(b'\x00', b'')
                    csv_text = clean_content.decode('utf-8', errors='replace')
                    csv_text = csv_text.strip()

                    try:
                        csv_file = io.StringIO(csv_text)
                        df = pd.read_csv(
                            csv_file, on_bad_lines='skip', skip_blank_lines=False)
                        content = f"CSV file with {len(df)} rows and {len(df.columns)} columns\n"
                        content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
                        for row_index, (_, row) in enumerate(df.iterrows()):
                            row_values = [
                                f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])]
                            if row_values:
                                content += f"Row {row_index + 1}: " + \
                                    " | ".join(row_values) + "\n"
                            else:
                                content += f"Row {row_index + 1}: [empty row]\n"
                        return content
                    except Exception:
                        return csv_text
            else:
                return "[Unsupported document tyoe]"
        except Exception as e:
            print(f"Error extracting full content: {str(e)}")
            return f"[Error extracting content: {str(e)}]"

    @staticmethod
    def _extract_segments(file_content: bytes, document_type: DocumentType, filename: str) -> tuple[Dict[str, Any], Dict[str, Any]]:
        """Extract segments and metadata based on document type"""
        try:
            if document_type == DocumentType.TEXT:
                segments_data, _ = DocumentUploadService._extract_text_content(
                    file_content)
                return segments_data, {}

            elif document_type == DocumentType.PDF:
                segments_data, _, file_metadata = DocumentUploadService._extract_pdf_content(
                    file_content)
                return segments_data, file_metadata

            elif document_type == DocumentType.DOCX:
                segments_data, _ = DocumentUploadService._extract_docx_content(
                    file_content)
                return segments_data, {}

            elif document_type == DocumentType.CSV:
                # Handle both CSV and Excel by converting spreadsheets to CSV text
                ext = pathlib.Path(filename or "").suffix.lower()
                if ext in (".xlsx", ".xls"):
                    # Convert Excel to CSV bytes
                    try:
                        excel_buffer = io.BytesIO(file_content)
                        df = pd.read_excel(excel_buffer)
                        csv_bytes = df.to_csv(index=False).encode('utf-8')
                    except Exception as e:
                        return {"segments": [], "error": str(e)}, {"error": str(e)}
                    segments_data, _, file_metadata = DocumentUploadService._extract_csv_content(
                        csv_bytes)
                else:
                    segments_data, _, file_metadata = DocumentUploadService._extract_csv_content(
                        file_content)
                return segments_data, file_metadata

            else:
                return {"segments": [], "error": "unsupported_type"}, {"error": "Unsupported document type"}

        except Exception as e:
            print(f"Error extracting segments: {str(e)}")
            return {"segments": [], "error": str(e)}, {"error": str(e)}

    @staticmethod
    def _extract_excel_content(file_content: bytes) -> tuple[Dict[str, Any], str, Dict[str, Any]]:
        """Extract content from Excel files with row-based segments for thematic analysis"""
        try:
            excel_file = io.BytesIO(file_content)
            df = pd.read_excel(excel_file)

            segments = []
            content_lines = []
            for row_index, (_, row) in enumerate(df.iterrows()):
                # Build text for each row
                values = [
                    f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])]
                if not values:
                    continue
                row_text = f"Row {row_index + 1}: " + " | ".join(values)
                content_lines.append(row_text)
                segments.append({
                    "type": "row",
                    "content": row_text,
                    "row_index": row_index,
                    "character_start": sum(len(line) + 2 for line in content_lines[:row_index]),
                    "character_end": sum(len(line) + 2 for line in content_lines[:row_index]) + len(row_text),
                    "additional_data": row.to_dict()
                })

            content = "\n\n".join(content_lines)
            structured_content = {
                "segments": segments,
                "total_segments": len(segments),
                "segmentation_type": "excel_row",
                "columns": df.columns.tolist()
            }
            metadata = {
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": df.columns.tolist(),
                "processing_note": "Converted Excel to row-based segments"
            }
            return structured_content, content, metadata
        except Exception as e:
            print(f"Error processing Excel file: {str(e)}")
            return {"segments": [], "error": str(e)}, "", {"error": str(e)}

    # @staticmethod
    # def _create_document_segments(db: Session, document: Document, segments_data: list):
    #     """Create DocumentSegment database records from extracted segment data using bulk insert"""
    #     segment_dicts = []
    #     for segment_data in segments_data:
    #         segment_dict = {
    #             "document_id": document.id,
    #             "segment_type": segment_data.get("type", "text"),
    #             "content": segment_data["content"],
    #             "line_number": segment_data.get("line_number"),
    #             "page_number": segment_data.get("page_number"),
    #             "paragraph_index": segment_data.get("paragraph_index"),
    #             "row_index": segment_data.get("row_index"),
    #             "character_start": segment_data.get("character_start"),
    #             "character_end": segment_data.get("character_end"),
    #             "additional_data": segment_data.get("additional_data"),
    #             "created_at": document.created_at,
    #             "updated_at": document.updated_at,
    #         }
    #         segment_dicts.append(segment_dict)
    #     # Bulk insert segments
    #     db.bulk_insert_mappings(DocumentSegment.__mapper__, segment_dicts)
    #     db.commit()
    #     print(f"Created {len(segments_data)} segments for document {document.id}")

    @staticmethod
    def _extract_text_content(file_content: bytes) -> tuple[Dict[str, Any], str]:
        """Extract content from text files with structured segments for thematic analysis"""
        try:
            # Clean the file content by removing NUL characters
            clean_content = file_content.replace(b'\x00', b'')

            # Split into lines and create segments
            text = clean_content.decode('utf-8', errors='replace')
            lines = text.split('\n')
            segments = []

            for i, line in enumerate(lines):
                line_content = line.strip()
                if line_content:
                    segments.append({
                        "type": "line",
                        "content": line_content,
                        "line_number": i + 1,
                        "character_start": sum(len(l) + 1 for l in lines[:i]),
                        "character_end": sum(len(l) + 1 for l in lines[:i]) + len(line)
                    })

            structured_content = {
                "segments": segments,
                "total_segments": len(segments),
                "segmentation_type": "line"
            }
            return structured_content, text

        except Exception as e:
            print(f"Error processing text file: {str(e)}")
            return {"segments": [], "error": str(e)}, f"[Error extracting text content: {str(e)}]"

    @staticmethod
    def _extract_pdf_content(file_content: bytes) -> tuple[Dict[str, Any], str, Dict[str, Any]]:
        """Extract content from PDF files with structured segments"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(pdf_file)

            content = ""
            all_segments = []
            segment_id = 1
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                content += f"\n--- Page {page_num + 1} ---\n"
                content += page_text
                content += "\n"

                # Split page text into lines for segmentation
                lines = page_text.split('\n')
                for line_index, line in enumerate(lines):
                    line_content = line.strip()
                    if line_content:
                        all_segments.append({
                            "type": "line",
                            "content": line_content,
                            "page_number": page_num + 1,
                            "line_number": line_index + 1,
                            "character_start": sum(len(l) + 1 for l in lines[:line_index]),
                            "character_end": sum(len(l) + 1 for l in lines[:line_index]) + len(line)
                        })
                        segment_id += 1

            structured_content = {
                "segments": all_segments,
                "total_segments": len(all_segments),
                "segmentation_type": "line_by_page"
            }
            metadata = {
                "page_count": len(pdf_reader.pages),
                "pdf_metadata": pdf_reader.metadata
            }

            return structured_content, content, metadata
        except Exception as e:
            return {"segments": [], "error": str(e)}, f"[Error extracting PDF content: {str(e)}]", {"error": str(e)}

    @staticmethod
    def _extract_docx_content(file_content: bytes) -> tuple[Dict[str, Any], str]:
        """Extract content from DOCX files with structured segments"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = DocxDocument(doc_file)

            content = ""
            segments = []
            segment_id = 1
            for paragraph_index, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:
                    content += para_text + "\n"

                    # Split paragraphs into sentences for better granularity
                    sentences = [s.strip()
                                 for s in para_text.split('.') if s.strip()]
                    for sentence_index, sentence in enumerate(sentences):
                        if sentence:
                            segments.append({
                                "type": "sentence",
                                "content": sentence + ("." if not sentence.endswith('.') else ""),
                                "paragraph_index": paragraph_index,
                                "character_start": len(content) - len(para_text) + sum(len(s) + 1 for s in sentences[:sentence_index]),
                                "character_end": len(content) - len(para_text) + sum(len(s) + 1 for s in sentences[:sentence_index + 1])
                            })
                            segment_id += 1

            structured_content = {
                "segments": segments,
                "total_segments": len(segments),
                "segmentation_type": "sentence"
            }
            return structured_content, content
        except Exception as e:
            return {"segments": [], "error": str(e)}, f"[Error extracting DOCX content: {str(e)}]"

    @staticmethod
    def _extract_csv_content(file_content: bytes) -> tuple[Dict[str, Any], str, Dict[str, Any]]:
        """Extract content from CSV files with row-based segments for thematic analysis"""
        try:
            # Clean the file content by removing NUL characters
            clean_content = file_content.replace(b'\x00', b'')

            # Decode with error handling and strip extra blank lines
            csv_text = clean_content.decode('utf-8', errors='replace').strip()
            csv_file = io.StringIO(csv_text)
            # Parse CSV rows, skipping irregular rows
            df = pd.read_csv(csv_file, on_bad_lines='skip',
                             skip_blank_lines=True)

            # Create segments for each row
            segments = []
            content_lines = []

            for row_index, (_, row) in enumerate(df.iterrows()):
                # Create a readable text entry for each row
                row_text = f"Row {row_index + 1}: "
                non_null_values = []

                for col, value in row.items():
                    if pd.notna(value):
                        non_null_values.append(f"{col}: {str(value)}")

                if non_null_values:
                    row_text += " | ".join(non_null_values)
                    content_lines.append(row_text)

                    segments.append({
                        "type": "row",
                        "content": row_text,
                        "row_index": row_index,
                        # +2 for double line breaks
                        "character_start": sum(len(line) + 2 for line in content_lines[:row_index]),
                        "character_end": sum(len(line) + 2 for line in content_lines[:row_index]) + len(row_text),
                        "additional_data": row.to_dict()
                    })

            # Join with double line breaks for readability
            content = "\n\n".join(content_lines)

            structured_content = {
                "segments": segments,
                "total_segments": len(segments),
                "segmentation_type": "csv_row",
                "columns": df.columns.tolist()
            }

            metadata = {
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": df.columns.tolist(),
                "processing_note": "Converted to row-based segments for thematic analysis"
            }
            return structured_content, "\n".join(content_lines), metadata

        except Exception as e:
            print(f"Error processing CSV file: {str(e)}")
            return {"segments": [], "error": str(e)}, f"[Error extracting CSV content: {str(e)}]", {"error": str(e)}
